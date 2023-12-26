from io import BytesIO
from flask_login import LoginManager, UserMixin, login_required, login_user, current_user, logout_user
from scipy.spatial.distance import euclidean
from scipy.stats import rankdata
from cmath import sqrt
from itertools import zip_longest
from MySQLdb import IntegrityError
from flask import Flask, render_template, request, send_file, url_for, redirect, flash, session
from flask_sqlalchemy import SQLAlchemy
import time
from datetime import datetime
import random
from collections import Counter
import pandas as pd
from docx import Document
import datetime


from flask_excel import make_response_from_query_sets, init_excel


app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'  # Use SQLite database
app.config['SECRET_KEY'] = 'your_secret_key'  # Change this to a secret key for session security
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'



class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    age = db.Column(db.Integer)
    gender = db.Column(db.String(10))
    major = db.Column(db.String(100))
    student_id = db.Column(db.String(20), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    pre_test_knowledge_gap = db.Column(db.Float, nullable=True)
    post_test_knowledge_gap = db.Column(db.Float, nullable=True)
    dominant_vark_style = db.Column(db.String(10), nullable=True)
    style_first = db.Column(db.String(10), nullable=True)
    style_second = db.Column(db.String(10), nullable=True)
    style_third = db.Column(db.String(10), nullable=True)
    style_fourth = db.Column(db.String(10), nullable=True)
    pre_test_duration = db.Column(db.String(10), nullable=True)

    def is_active(self):
        return True
        

    def __repr__(self):
        return f"User('{self.name}', '{self.email}', '{self.student_id}')"


class Performance(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.ForeignKey('user.id'), nullable=False)
    pre_test = db.Column(db.Text, nullable=True)
    post_test = db.Column(db.Text, nullable=True)

    # Define a relationship with the User model
    user = db.relationship('User', backref=db.backref('performances', lazy=True))

    def __repr__(self):
        return f"Performance(id={self.id}, user_id={self.user_id}, pre_test={self.pre_test}, post_test={self.post_test})"


class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    question_text = db.Column(db.String(500), nullable=False)
    options = db.relationship('Option', backref='question', lazy=True)


class Option(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    option_id = db.Column(db.String(1), nullable=False)
    content = db.Column(db.String(500), nullable=False)
    

class QuestionnaireResponse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    answer_option_id = db.Column(db.String(1), nullable=False)
    vark_score = db.Column(db.Float, nullable=True)
    knowledge_gap_score = db.Column(db.Float, nullable=True)
    time_spent_score = db.Column(db.Float, nullable=True)
    performance_score = db.Column(db.Float, nullable=True)

    user = db.relationship('User', backref=db.backref('questionnaire_responses', lazy=True))
    question = db.relationship('Question', backref=db.backref('responses', lazy=True))

class QuestionMC(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    question_text = db.Column(db.String(500), nullable=False)
    options = db.relationship('OptionMC', backref='question', lazy=True)
    category = db.Column(db.String, nullable=True)


class OptionMC(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    
    question_id = db.Column(db.Integer, db.ForeignKey('question_mc.id'), nullable=False)
    option_id = db.Column(db.String(1), nullable=False) 
    content = db.Column(db.String(500), nullable=False)



class UserResponse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    question_number = db.Column(db.Integer, nullable=False)
    user_response = db.Column(db.String(1), nullable=False)




# Create tables before running the app
with app.app_context():
    db.create_all()



@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def is_logged_in():
    return 'user_id' in session

# Define routes for the admin pagesslow jams
@app.route('/admin_page')
def admin_dashboard():
    return render_template('admin.html')

@app.route('/view_users')
def view_users():
    # Replace the following line with the logic to get the list of users
    users = User.query.all()
    return render_template('view_users.html', users=users)

import xlsxwriter
# Route to download users data in Excel format
@app.route('/admin/download_excel')
def download_excel():
    # Fetch all users from the database
    users = User.query.all()

    # Create an in-memory Excel file
    output = BytesIO()

    # Create a workbook and add a worksheet
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet()

    # Add header row
    header_row = ['ID', 'Name', 'Email', 'Age', 'Gender', 'Major', 'Student ID',
                  'Pre-test Knowledge Gap', 'Post-test Knowledge Gap',
                  'Dominant VARK Style', 'Style First', 'Style Second', 'Style Third', 'Style Fourth']
    for col_num, header in enumerate(header_row):
        worksheet.write(0, col_num, header)

    # Add user data rows
    for row_num, user in enumerate(users, start=1):
        worksheet.write(row_num, 0, user.id)
        worksheet.write(row_num, 1, user.name)
        worksheet.write(row_num, 2, user.email)
        worksheet.write(row_num, 3, user.age)
        worksheet.write(row_num, 4, user.gender)
        worksheet.write(row_num, 5, user.major)
        worksheet.write(row_num, 6, user.student_id)
        worksheet.write(row_num, 7, user.pre_test_knowledge_gap)
        worksheet.write(row_num, 8, user.post_test_knowledge_gap)
        worksheet.write(row_num, 9, user.dominant_vark_style)
        worksheet.write(row_num, 10, user.style_first)
        worksheet.write(row_num, 11, user.style_second)
        worksheet.write(row_num, 12, user.style_third)
        worksheet.write(row_num, 13, user.style_fourth)

    # Close the workbook
    workbook.close()

    # Set the position to the beginning of the stream
    output.seek(0)

    # Serve the file for download
    return send_file(output, as_attachment=True, download_name='users.xlsx')


###############################################################
#####                                                     #####
#####                     HOME PAGE                       #####
#####                                                     #####
############################################################### 

# Step 1: User Registration and Profile Creation
@app.route('/')
def home():
    return render_template('home.html')





###############################################################
#####                                                     #####
#####                      Register                       #####
#####                                                     #####
############################################################### 

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    
    if request.method == 'POST':

        user_register = [None, None, None, None, None, None, None]

        action = request.form.get("action")
        email = request.form.get('email')
        name = request.form.get('username')
        student_id = request.form.get('student_id')
        age = request.form.get('age')
        gender = request.form.get('gender') #
        major = request.form.get('major') #
        password = request.form.get('password')
        password_2 = request.form.get('passwaord_2')

        if action == "Continue":

            existing_user = User.query.filter_by(email=email).first()

            if existing_user:
                flash('Username already exists. Please choose another username.', 'danger')
                return render_template('Register/signup.1.html')
            
            else:
                user_register[0] = email
                print(user_register)
                return render_template('Register/signup.2.html' , Email = email)
            
        if action == "Continue 2":
            studentID= User.query.filter_by(student_id=student_id).first()

            if studentID:
                flash('Email with Student ID already exists. Please choose another Student ID.', 'danger')
                return render_template('Register/signup.2.html')
            
            else:
                user_register[0] = email
                user_register[1] = name
                user_register[2] = student_id
                print(user_register)
                return render_template('Register/signup.3.html' , Email = email, Name=name, Student_ID=student_id)
        
        if action == "Continue 3":
            user_register[0] = email
            user_register[1] = name
            user_register[2] = student_id
            user_register[3] = age
            user_register[4] = major
            user_register[5] = gender
            print(user_register)
            return render_template('Register/signup.4.html' , Email = email, Name=name, Student_ID=student_id, Age=age, Major=major, Gender=gender)
        

        if action== "Sign up":
            user_register[0] = email
            user_register[1] = name
            user_register[2] = student_id
            user_register[3] = age
            user_register[4] = major
            user_register[5] = gender
            user_register[6] = password
            new_user = User(name=user_register[1], email=user_register[0], age=user_register[3], gender=user_register[5], student_id=user_register[2], password=password)
            try:
                print(user_register)
                db.session.add(new_user)
                db.session.commit()
                flash('Account created successfully. Please log in.', 'success')
                return redirect(url_for('login'))

            except IntegrityError:
                db.session.rollback()
                flash('Username already exists. Please choose another username.', 'danger')
                return render_template('signup.2.html, Email = email')


    return render_template('Register/signup.1.html')



###############################################################
#####                                                     #####
#####                        LOGIN                        #####
#####                                                     #####
############################################################### 

# Step 2: User Login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
     
    if request.method == 'POST':
        email = request.form.get('username')
        password = request.form.get('password')

        # Check if the user exists and the password is correct
        user = User.query.filter_by(email=email).first()
        if user and user.password == password:
            login_user(user)
            flash('Login successful!', 'success')
            session['user_id'] = user.id
            if "CreateAccount" not in session:
                session["CreateAccount"] = datetime.datetime.now()
                
            return redirect(url_for('home'))
        else:
            flash('Login failed. Check your student ID and password.', 'danger')

    return render_template('login.html')



###############################################################
#####                                                     #####
#####                        LOGOUT                       #####
#####                                                     #####
############################################################### 

@app.route('/logout')
def logout():
    if is_logged_in():
        session.pop('user_id', None)
        logout_user()
        flash('You have been logged out.', 'success')
    else:
        flash('You are not currently logged in.', 'info')


    return redirect(url_for('login'))


###############################################################
#####                                                     #####
#####                  User Profile                       #####
#####                                                     #####
############################################################### 

@app.route('/profile')
def user_profile():
    # Check if the user is logged in
    if not is_logged_in():
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))

    # Retrieve user information from the database
    user_id = session['user_id']
    user = User.query.get(user_id)

    # # Calculate the dominant VARK learning style for the user
    # dominant_vark = calculate_vark_profile(user_id)
    # performance_scores = calculate_performance_scores(user_id)
    question_mc_responses = QuestionMC.query.all()
    user_responses = UserResponse.query.filter_by(user_id=user_id).all()
    # List of correct answers in order
    correct_answers = ['C', 'B', 'A', 'B', 'C', 'B', 'B', 'C', 'C', 'A', 'B', 'B', 'C', 'B', 'B', 'B', 'B', 'A', 'B', 'B']
    # Check if the user has completed the post-test
    pre_test_duration = 0.0
    post_test_duration = 0.0
    created_Account = ""
    total_time_taken = "Test not Finished"

    if "pre_test_start_time" in session and "pre_test_end_time" in session:
        start_time = session["pre_test_start_time"]
        end_time = session["pre_test_end_time"]
        start_time = start_time.replace(tzinfo=None)
        end_time = end_time.replace(tzinfo=None)

        pre_test_duration = (end_time - start_time).total_seconds()
        pre_test_duration = "{:.2f}".format(pre_test_duration/60)
    
    if "post_test_start_time" in session and "post_test_end_time" in session:
        start_time = session["post_test_start_time"]
        end_time = session["post_test_end_time"]
        start_time = start_time.replace(tzinfo=None)
        end_time = end_time.replace(tzinfo=None)

        post_test_duration = (end_time - start_time).total_seconds()
        post_test_duration = "{:.2f}".format(post_test_duration/60)

        created_Account = session["CreateAccount"]
        created_Account = created_Account.replace(tzinfo=None)
        total_time_taken = "{:.2f}".format(created_Account/60)


    if "CreateAccount" in session:
        created_Account = session["CreateAccount"]
                
    post_test_completed = False

    if user.post_test_knowledge_gap is not None:
        post_test_completed = True
    # Pass the additional information to the template
    return render_template('profile2.html', user=user, question_mc_responses=question_mc_responses, user_responses=user_responses, get_style_name=get_style_name, correct_answers=correct_answers, post_test_completed=post_test_completed, pre_test_duration=pre_test_duration, post_test_duration=post_test_duration, created_Account=created_Account,total_time_taken=total_time_taken)



@app.route('/update_profile', methods=['GET', 'POST'])
@login_required  # Ensure the user is logged in to access this page
def update_profile():

    if request.method == 'POST':
        # Handle the form submission
        fname = request.form.get('fname')
        lname = request.form.get('lname')
        phone = request.form.get('phone')

        # Update the user's information
        current_user.name = f"{fname} {lname}"
        # current_user.phone = phone

        # Commit changes to the database
        db.session.commit()

        flash('Profile updated successfully', 'success')

        return redirect(url_for('user_profile'))

    return render_template('user_profile.html', user=current_user)

@app.route('/download_profile_docx')
def download_profile_docx():
    # Check if the user is logged in
    if not is_logged_in():
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))

    # Retrieve user information from the database
    user_id = session['user_id']
    user = User.query.get(user_id)

    # Retrieve pre-test quiz questions and user responses
    question_mc_responses = QuestionMC.query.all()
    user_responses = UserResponse.query.filter_by(user_id=user_id).all()

    # List of correct answers in order
    correct_answers = ['C', 'B', 'A', 'B', 'C', 'B', 'B', 'C', 'C', 'A', 'B', 'B', 'C', 'B', 'B', 'B', 'B', 'A', 'B', 'B']

    # Create a Word document
    doc = Document()
    doc.add_heading("FlexLearning-MCDM", level=0)

    doc.add_heading('User Profile and Pre-test Results', level=1)

    # Add user information
    doc.add_heading('User Information', level=2)
    doc.add_paragraph(f'Name: {user.name}')
    doc.add_paragraph(f'Email: {user.email}')
    doc.add_paragraph(f'Dominant Learning Style: {get_style_name(user.dominant_vark_style)}')


    doc.add_heading('Pre-test Quizzes and Responses', level=1)

    for i, question in enumerate(question_mc_responses):
        doc.add_heading(f'Question {i + 1}', level=2)
        doc.add_paragraph(f'Question: {question.question_text}')
        doc.add_paragraph(f'Category: {question.category}')
        doc.add_paragraph('Options:')
        for option in question.options:
            doc.add_paragraph(f'{option.option_id}. {option.content}')
        doc.add_paragraph(f'Your Response: {next((response.user_response for response in user_responses if response.question_number == question.id), "N/A")}')
        doc.add_paragraph(f'Correct Answer: {correct_answers[i]}')
        doc.add_paragraph('')  # Add an empty paragraph for better readability

    # Save the document to a temporary file
    docx_filename = 'pretest_responses.docx'
    doc.save(docx_filename)

    # Serve the file for download
    return send_file(docx_filename, as_attachment=True, download_name='pretest_responses.docx')

import time
# step  vark questionnares
@app.route('/questionnaire/<int:question_number>', methods=['GET', 'POST'])
def vark_questionnaire(question_number):
    if not current_user.is_authenticated:
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))

    
    if current_user.dominant_vark_style:
        # Redirect to a page indicating that the user needs to complete the VARK questionnaire first
        flash("Already handled the vark questionares", "danger")
        return redirect(url_for('Pre_Test_Quiz', question_number=1))
    
     # Retrieve the start time from the session when the questionnaire begins
    if 'questionnaire_start_time' not in session:
        session['questionnaire_start_time'] = time.time()

     # Retrieve the question and options for the specified question number
    current_question = Question.query.filter_by(id=question_number).first()

    if current_question is None:
        user_id = session['user_id']
        user = User.query.get(user_id)

        vark_responses = QuestionnaireResponse.query.all()


        # Create a DataFrame from the responses
        columns = ['User ID', 'Question Number', 'Answer Option ID', 'VARK Score', 'Knowledge Gap Score', 'Time Spent Score', 'Performance Score']
        data = [(response.user_id, response.question_id, response.answer_option_id,
                response.vark_score, response.knowledge_gap_score, response.time_spent_score, response.performance_score)
                for response in vark_responses]
        df = pd.DataFrame(data, columns=columns)


        # Export DataFrame to Excel
        file_path = 'vark_responses.xlsx'
        df.to_excel(file_path, index=False)


        # Calculate performance scores for all responses
        all_responses = QuestionnaireResponse.query.filter_by(user_id=user_id).all()
        performance_scores = [[response.vark_score, response.knowledge_gap_score, response.time_spent_score] for response in all_responses]

        #  Calculate total time spent
        questionnaire_start_time = session['questionnaire_start_time']
        total_time_spent = time.time() - questionnaire_start_time


        # Get TOPSIS ranking
        topsis_ranks = topsis_ranking(performance_scores)


        # Calculate the dominant VARK learning style for the user
        dominant_vark = calculate_vark_profile(user_id)


        user_vark = User.query.get(current_user.id)
        user_vark.dominant_vark_style = dominant_vark
        db.session.commit()


        # performance_scores = calculate_performance_scores(user_id)
        vark_recommendation, topsis_recommendation, ranked_vark= generate_recommendation(user_id)
        vark_material_recommendation = recommend_materials(dominant_vark)
        

        user_vark = User.query.get(current_user.id)
        user_vark.style_first = ranked_vark[0]
        user_vark.style_second = ranked_vark[1]
        user_vark.style_third = ranked_vark[2]
        user_vark.style_fourth = ranked_vark[3]
        db.session.commit()


        # Combine questions and ranks for rendering
        questions_ranks = list(zip_longest(Question.query.all(), topsis_ranks))

        return render_template('profile2.html',user=user,ranked_vark=ranked_vark, dominant_vark=dominant_vark, performance_scores=performance_scores, questions_ranks=questions_ranks, total_time_spent=total_time_spent, vark_recommendation=vark_recommendation, topsis_recommendation=topsis_recommendation, vark_material_recommendation=vark_material_recommendation, get_style_name=get_style_name)
    
    
    if request.method == 'POST':
        # Increment the question number for the next iteration
        if request.form.get('answer_option_id') is not None:
            action = request.form.get('action')

            if action == "next":
                next_question_number = question_number + 1
                

            elif action == "back" and question_number > 1:
                next_question_number = question_number - 1

            
            # Calculate scores based on some hypothetical criteria
            answer_option_id = request.form.get('answer_option_id')
            user_id = session['user_id']


            vark_score = calculate_vark_score(answer_option_id)
            knowledge_gap_score = calculate_knowledge_gap_score()
            time_spent = random.uniform(0.4, 1)
            time_spent_score = time_spent
            performance_score = calculate_performance_score(vark_score, knowledge_gap_score, time_spent_score)

            # Create QuestionnaireResponse instance with calculated scores
            response = QuestionnaireResponse(
                user_id=user_id,
                question_id=question_number,
                answer_option_id=answer_option_id,
                vark_score=vark_score,
                knowledge_gap_score=knowledge_gap_score,
                time_spent_score=time_spent_score,
                performance_score=performance_score
            )


            # response = QuestionnaireResponse(user_id = user_id, question_id=question_number, answer_option_id=answer_option_id)
            max_allowed_time = 1800
            db.session.add(response)
            db.session.commit()

            return redirect(url_for('vark_questionnaire', question_number=next_question_number))
        
        else:
            flash('Click on a choice.', 'info')

            

    # Fetch options for the current question
    options = Option.query.filter_by(question_id=question_number).all()

    return render_template('vark_questionnaire.html', question=current_question.question_text, options=options, question_number=question_number, T_num_Q=Question.query.count())

# Example function to map style codes to names
def get_style_name(style_code):
    style_mapping = {
        'A': 'Audio',
        'V': 'Visual',
        'K': 'Kinesthetic',
        'R': 'Read/Write',
        # Add more styles as needed
    }
    return style_mapping.get(style_code, 'Unknown Style')


vark_test_response = {}

###############################################################
#####                                                     #####
#####                  Pre Test                           #####
#####                                                     #####
############################################################### 

@app.route('/Pre_testQuiz/<int:question_number>', methods=['GET', 'POST'])
@login_required
def Pre_Test_Quiz(question_number):
    if not is_logged_in():
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))
    
    # Check if the user has completed the VARK questionnaire
    if not current_user.dominant_vark_style:
        # Redirect to a page indicating that the user needs to complete the VARK questionnaire first
        flash("You have to handle the vark questionnare first", "danger")
        return redirect(url_for('vark_questionnaire', question_number=1))
    
    if current_user.pre_test_knowledge_gap:
        # Redirect to a page indicating that the user needs to complete the VARK questionnaire first
        flash("You have Already handled the Pre-Test questionnare", "danger")
        return redirect(url_for('Post_Test_Quiz', question_number=1))
   

    current_question = QuestionMC.query.filter_by(id=question_number).first()
    if question_number == 1:
        if "pre_test_start_time" not in session:
            session["pre_test_start_time"] = datetime.datetime.now()

    
    if current_question is None or current_question == 20:
        knowledge_gap = calculate_knowledge_gap()

        if "pre_test_start_time" in session:
            session["pre_test_end_time"] = datetime.datetime.now()
                        # duration = (session["pre_test_end_time"] - session["pre_test_start_time"]).total_seconds()
        return render_template('thank_you.html', knowledge_gap=knowledge_gap)
    
    
    if request.method == 'POST':
        # Increment the question number for the next iteration
        user_response = request.form.get('answer_option_id')
        if request.form.get('answer_option_id') is not None:
            action = request.form.get('action')

            # Store user response in the database
            response = UserResponse(question_number=question_number, user_response=user_response, user_id=current_user.id)
            db.session.add(response)
            db.session.commit()

            if action == "next":
                next_question_number = question_number + 1 
                if question_number == 20:
                    user_responses = UserResponse.query.all()
                    
                                    # Get pre-test data from the database
                    pre_test_responses = UserResponse.query.all()

                   
                    pre_test_knowledge_gap = calculate_knowledge_gap()
                    user = User.query.get(current_user.id)
                    user.pre_test_knowledge_gap = pre_test_knowledge_gap
                    db.session.commit()

                    if "pre_test_start_time" in session:
                        session["pre_test_end_time"] = datetime.datetime.now()
                        # duration = (session["pre_test_end_time"] - session["pre_test_start_time"]).total_seconds()

            
            elif action == "back" and question_number > 1:
                next_question_number = question_number - 1
            
            return redirect(url_for('Pre_Test_Quiz', question_number=next_question_number))
        
        else:
            flash('Click on a choice.', 'info')

    options = OptionMC.query.filter_by(question_id=question_number).all()
    return render_template('testQuizes.html', question=current_question.question_text, options=options, question_number=question_number, T_num_Q=20)
 


###############################################################
#####                                                     #####
#####                  post Test                           #####
#####                                                     #####
############################################################### 

@app.route('/postTestQuiz/<int:question_number>', methods=['GET', 'POST'])
@login_required
def Post_Test_Quiz(question_number):
    if not is_logged_in():
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))
    
    # Check if the user has completed the VARK questionnaire
    if not current_user.dominant_vark_style:
        flash("You have to handle the vark questionnare first", "danger")
        return redirect(url_for('vark_questionnaire', question_number=1))
    
    if not current_user.pre_test_knowledge_gap:
        # Redirect to a page indicating that the user needs to complete the VARK questionnaire first
        flash("You have to handle the Pre-Test questionnare first", "danger")
        return redirect(url_for('Pre_Test_Quiz', question_number=1))
   
    current_question = QuestionMC.query.filter_by(id=question_number).first()
    if question_number == 1:
        if "post_test_start_time" not in session:
            session["post_test_start_time"] = datetime.datetime.now()
    


    if current_question is None or current_question == 20:
        knowledge_gap = calculate_knowledge_gap()
        if "post_test_start_time" in session:
            session["post_test_end_time"] = datetime.datetime.now()
                        # duration = (session["pre_test_end_time"] - session["pre_test_start_time"]).total_seconds()
        return render_template('thank_you.html', knowledge_gap=knowledge_gap)
    
    
    if request.method == 'POST':
        # Increment the question number for the next iteration
        user_response = request.form.get('answer_option_id')
        if request.form.get('answer_option_id') is not None:
            action = request.form.get('action')

            # Store user response in the database
            response = UserResponse(question_number=question_number, user_response=user_response, user_id=current_user.id)
            db.session.add(response)
            db.session.commit()

            if action == "next":
                next_question_number = question_number + 1 
                if question_number == 20:
                    post_test_knowledge_gap = calculate_knowledge_gap()
                    user = User.query.get(current_user.id)
                    user.post_test_knowledge_gap = post_test_knowledge_gap
                    db.session.commit()
                    return render_template('thank_you.html', knowledge_gap=post_test_knowledge_gap)
            
            elif action == "back" and question_number > 1:
                next_question_number = question_number - 1
            
            return redirect(url_for('Post_Test_Quiz', question_number=next_question_number))
        
        else:
            flash('Click on a choice.', 'info')

    options = OptionMC.query.filter_by(question_id=question_number).all()
    return render_template('postTestQuiz.html', question=current_question.question_text, options=options, question_number=question_number, T_num_Q=20)
 


def calculate_knowledge_gap():
    # Retrieve all user responses from the database
    user_responses = UserResponse.query.all()

    total_questions = 20 # Assuming you have 20 questions

    # Initialize counters for correct and total responses
    correct_responses = 0
    total_responses = 0
    print("calculating user knowledge gap")

    print("calculate user reponse")

    for response in user_responses:
        total_responses += 1
        print(total_responses)

        # Get the correct answer for the question
        correct_answer = correct_answer_for_question(response.question_number)

        # Check if the user response matches the correct answer
        if response.user_response == correct_answer:
            # print(response.user_response)
            print(response.user_response)
            print(correct_answer)
            correct_responses += 1
        
        print(response.user_response)
        print(correct_answer)
    # Calculate knowledge gap as the percentage of correct responses
    if total_responses > 0:
        knowledge_gap_percentage = (correct_responses / total_responses) * 100
        print(knowledge_gap_percentage)
    else:
        knowledge_gap_percentage = 0  # Handle the case where there are no responses

    return knowledge_gap_percentage


# List of correct answers in order
correct_answers = ['C', 'B', 'A', 'B', 'C', 'B', 'B', 'C', 'C', 'A', 'B', 'B', 'C', 'B', 'B', 'B', 'B', 'A', 'B', 'B']

def correct_answer_for_question(question_number):
    # Check if the question_number is within the valid range
    if 1 <= question_number <= len(correct_answers):
        # Adjust index since question_number starts from 1
        return correct_answers[question_number - 1]
    else:
        # Handle the case where the question_number is out of range
        return None  # You might want to return a default value or raise an exception



@app.route('/recommendations')
@login_required
def recommendations():
    if not is_logged_in():
        flash('You need to log in first.', 'warning')
        return redirect(url_for('login'))
    
    # Check if the user has completed the VARK questionnaire
    if not current_user.dominant_vark_style:
        flash("You have to handle the vark questionnaire first", "danger")
        return redirect(url_for('vark_questionnaire', question_number=1))
    
    if not current_user.pre_test_knowledge_gap:
        # Redirect to a page indicating that the user needs to complete the VARK questionnaire first
        flash("You have to handle the Pre-Test questionnare first", "danger")
        return redirect(url_for('Pre_Test_Quiz', question_number=1))
    
    # Assuming 'pretest_results' is a field in the User model representing pretest scores
    pretest_results = current_user.pre_test_knowledge_gap

    # Get user's dominant learning style
    dominant_learning_style = current_user.dominant_vark_style

    
    # Get recommendations based on the user's dominant learning style and pretest results
    all_recommendations = get_all_recommendations(dominant_learning_style, pretest_results)

    # Separate recommendations into categories
    audio_recommendations = [rec for rec in all_recommendations if rec['type'] == 'audio']
    video_recommendations = [rec for rec in all_recommendations if rec['type'] == 'video']
    writing_recommendations = [rec for rec in all_recommendations if rec['type'] == 'writing']
    correct_answers = ['C', 'B', 'A', 'B', 'C', 'B', 'B', 'C', 'C', 'A', 'B', 'B', 'C', 'B', 'B', 'B', 'B', 'A', 'B', 'B']

    return render_template('recommendations.html', user=current_user,
                           audio_recommendations=audio_recommendations,
                           video_recommendations=video_recommendations,
                           writing_recommendations=writing_recommendations, dominant_learning_style=dominant_learning_style)
# Inside your Flask app
def get_all_recommendations(dominant_learning_style, pretest_results):
    recommendations = []

    # Logic for audio recommendations
    if dominant_learning_style == 'A':
        recommendations.append({'type': 'audio', 'title': 'Audio Title 1', 'description': 'Description 1', 'link': 'audio_link_1'})
        recommendations.append({'type': 'audio', 'title': 'Audio Title 2', 'description': 'Description 2', 'link': 'audio_link_2'})

    # Logic for video recommendations
    elif dominant_learning_style == 'V':
        recommendations.append({'type': 'video', 'title': 'Video Title 1', 'description': 'Description 1', 'link': 'video_link_1'})
        recommendations.append({'type': 'video', 'title': 'Video Title 2', 'description': 'Description 2', 'link': 'video_link_2'})

    # Logic for writing recommendations
    elif dominant_learning_style == 'R':
        recommendations.append({'type': 'writing', 'title': 'Writing Title 1', 'description': 'Description 1', 'link': 'writing_link_1'})
        recommendations.append({'type': 'writing', 'title': 'Writing Title 2', 'description': 'Description 2', 'link': 'writing_link_2'})

    # Common logic or additional logic based on pretest results can be added here

    return recommendations



# Define a route for the topics page
@app.route('/topics')
@login_required
def topics():
    if not current_user.dominant_vark_style:
        flash("You have to handle the vark questionnare first", "danger")
        return redirect(url_for('vark_questionnaire', question_number=1))
    # List of topics
    topics = {
        "Storage Components": "Components responsible for storing data in a computer system.",
        "Central Processing Unit (CPU)": "The core processing unit responsible for executing instructions.",
        "Motherboard and Communication Management": "Main circuit board managing communication between components.",
        "Primary Input Devices": "Devices used for input, such as keyboards and mice.",
        "Motherboard Functions": "Functions and roles of the motherboard in a computer system.",
        "Types of Software": "Different categories of software enabling various functionalities.",
        "Operating System (OS) Purpose": "The role and purpose of an operating system in a computer.",
        "Security Software": "Software designed to protect a computer system from security threats.",
        "Programming Languages": "Languages used for writing software and applications.",
        "Database Management Software": "Software for managing and organizing databases.",
        "RAM Functions": "Functions of Random Access Memory (RAM) in a computer.",
        "Types of Memory (RAM, ROM, Cache, Virtual)": "Different types of memory in a computer system.",
        "Non-volatile Memory": "Memory that retains data even when the power is off.",
        "Cache Memory Purpose": "The purpose and role of cache memory in a computer.",
        "Memory Speed Enhancement": "Methods for enhancing the speed of memory in a computer.",
        "Types of Ports": "Different types of ports for connecting devices to a computer.",
        "Video Output Ports": "Ports specifically used for video output.",
        "Networking Ports (HTTP, Ethernet)": "Ports related to networking and data transfer protocols.",
        "Ethernet Port Function": "Functionality and purpose of Ethernet ports.",
        "HDMI Port Purpose": "The role and purpose of HDMI ports in connecting audio and video devices."
    }

    # Render the template and pass the list of topics to it
    return render_template('topics.html', topics=topics)

@app.route('/re_take_test')
def ReTakeTest():
    user = User.query.get(current_user.id)
    user.post_test_knowledge_gap = None
    user.pre_test_knowledge_gap = None
    user.dominant_vark_style = None
    db.session.commit()

    return redirect(url_for('vark_questionnaire', question_number=1))

def get_recommendations_based_on_learning_style(dominant_learning_style):
    # This is a placeholder. Replace it with your actual recommendation logic
    # The recommendations should be filtered based on the user's dominant learning style
    if dominant_learning_style == 'V':
        return [
            {"title": "Visual Learning Course 1", "description": "Courses with visual content to enhance learning."},
            {"title": "Visual Learning Course 2", "description": "Another course with visual materials."},
            {"title": "Interactive Visual Presentations", "description": "Engaging presentations with graphical elements."},
            {"title": "Visual Design Principles Workshop", "description": "Learn principles of effective visual design."},
            {"title": "Data Visualization Masterclass", "description": "Explore techniques for impactful data visualization."},
        ]
    elif dominant_learning_style == 'A':
        return [
            {"title": "Auditory Learning Series 1", "description": "Audio-based learning resources for better retention."},
            {"title": "Auditory Learning Series 2", "description": "Another series with auditory content."},
            {"title": "Educational Podcasts Compilation", "description": "Diverse podcasts covering various educational topics."},
            {"title": "Language Learning with Audio", "description": "Improve language skills through audio-based lessons."},
            {"title": "Audiobooks Collection", "description": "Listen to books on a wide range of subjects."},
        ]
    elif dominant_learning_style == 'R':
        return [
            {"title": "Reading-focused Materials 1", "description": "Materials suited for reading-based learners."},
            {"title": "Reading-focused Materials 2", "description": "Additional resources for those who prefer reading."},
            {"title": "E-books Library", "description": "Access a variety of e-books on different subjects."},
            {"title": "Online Articles Repository", "description": "Curated collection of informative online articles."},
            {"title": "Reading Comprehension Exercises", "description": "Enhance reading skills with comprehension exercises."},
        ]
    elif dominant_learning_style == 'K':
        return [
            {"title": "Kinesthetic Learning Activities 1", "description": "Hands-on activities for kinesthetic learners."},
            {"title": "Kinesthetic Learning Activities 2", "description": "More activities to engage kinesthetic learners."},
            {"title": "Physical Education Tutorials", "description": "Learn through physical movement and exercise."},
            {"title": "DIY Science Experiments", "description": "Perform experiments for a hands-on learning experience."},
            {"title": "Art and Craft Workshops", "description": "Express creativity through kinesthetic arts and crafts."},
        ]

    # Default recommendations for unknown or unsupported learning styles
    return [
        {"title": "General Learning Resources", "description": "Explore a variety of learning materials suitable for all styles."}
    ]

VARK_WEIGHTS = {
    'V': 0.25,
    'A': 0.25,
    'R': 0.25,
    'K': 0.25,
}

CRITERIA_WEIGHTS = {
    'VARK': 0.5,
    'KnowledgeGap': 0.3,
    'TimeSpent': 0.2,
}

def calculate_vark_profile(user_id):
    # Retrieve user's responses from the database
    user_responses = QuestionnaireResponse.query.filter_by(user_id=user_id).all()

    # Count occurrences of each VARK category
    # vark_counts = {'V': 0, 'A': 0, 'R': 0, 'K': 0}
    vark_counts = Counter()

    for response in user_responses:
        vark_counts[response.answer_option_id] += 1

    # Determine dominant VARK category
    dominant_vark = max(vark_counts, key=vark_counts.get)

    return dominant_vark




def normalize_values(criterion_values):
    max_value = max(criterion_values)
    min_value = min(criterion_values)

    normalized_values = [(value - min_value) / (max_value - min_value) for value in criterion_values]

    return normalized_values


def assign_weights(criteria_values, weights):
    weighted_values = [value * weights[index] for index, value in enumerate(criteria_values)]

    return weighted_values



# calculate_performance_scores function to calculate scores for each criterion and the overall performance score.
def calculate_performance_scores(user_id):
    user_responses = QuestionnaireResponse.query.filter_by(user_id=user_id).all()
    
    for response in user_responses:
        response.vark_score = VARK_WEIGHTS.get(response.answer_option_id, 0.0)
        response.knowledge_gap_score = response.knowledge_gap_score if response.knowledge_gap_score is not None else 0.0
        response.time_spent_score = response.time_spent_score if response.time_spent_score is not None else 0.0

        # Calculate performance score
        response.performance_score = (
            response.vark_score * CRITERIA_WEIGHTS['VARK'] +
            response.knowledge_gap_score * CRITERIA_WEIGHTS['KnowledgeGap'] +
            response.time_spent_score * CRITERIA_WEIGHTS['TimeSpent']
        )

    db.session.commit()




# Create a function to implement the TOPSIS method to rank the responses based on their performance scores.

def rank_content_items(scores):
    # Convert complex numbers to floats
    scores = [float(score.real) if isinstance(score, complex) else score for score in scores]

    # Handle None values by replacing them with a large negative value
    scores = [float('-inf') if score is None else score for score in scores]

    # Print the scores before sorting
    print("Scores before sorting:", scores)

    # Sort content items based on scores
    ranked_content = sorted(enumerate(scores, 1), key=lambda x: x[1], reverse=True)

    # Print the sorted content for debugging
    print("Sorted content:", ranked_content)

    # Extract the ranked indices
    ranked_indices = [index for index, _ in ranked_content]

    return ranked_indices



def topsis_ranking(performance_scores):
    # Transpose the performance scores matrix
    transposed_scores = list(map(list, zip(*performance_scores)))

    # Calculate squared sums of positive and negative values
    positive_squared_sums = [sum([value**2 for value in column if value > 0]) for column in transposed_scores]
    negative_squared_sums = [sum([value**2 for value in column if value < 0]) for column in transposed_scores]

    # Calculate square roots
    positive_square_roots = [sqrt(value) for value in positive_squared_sums]
    negative_square_roots = [sqrt(value) for value in negative_squared_sums]

    # Calculate TOPSIS scores
    topsis_scores = [positive / (positive + negative) if (positive and negative) else None for positive, negative in zip(positive_square_roots, negative_square_roots)]

    # Print the values before ranking
    print("TOPSIS Scores:", topsis_scores)

    # Rank the questions based on TOPSIS scores
    topsis_ranks = rank_content_items(topsis_scores)
    print(topsis_ranks)

    return topsis_ranks



def calculate_vark_score(answer_option_id):
    """
    Calculate VARK score based on the selected answer option.
    This is just a hypothetical example, replace it with your actual logic.
    """
    if answer_option_id == 'V':
        return 0.8
    elif answer_option_id == 'A':
        return 0.6
    elif answer_option_id == 'R':
        return 0.4
    elif answer_option_id == 'K':
        return 0.7
    else:
        return 0.0
    
    
def calculate_knowledge_gap_score():
    
    return random.uniform(0.1, 1)



def calculate_time_spent_score(total_time_spent, max_allowed_time):
    """
    Calculate time spent score based on the total time spent on the questionnaire.
    Replace this with your actual logic.
    """
    # Assuming max_allowed_time is the maximum time allowed for the questionnaire
    normalized_time_spent = total_time_spent / max_allowed_time

    # You might want to adjust the weights and conditions based on your specific requirements
    if normalized_time_spent <= 0.5:
        return 1.0  # Full score for efficient completion
    elif 0.5 < normalized_time_spent <= 0.75:
        return 0.7  # Partial score for moderate completion time
    else:
        return 0.5  # Lower score for longer completion time

def calculate_performance_score(vark_score, knowledge_gap_score, time_spent_score):
    """
    Calculate overall performance score based on individual scores.
    Replace this with your actual logic.
    """
    return 0.5 * vark_score + 0.3 * knowledge_gap_score + 0.2 * time_spent_score

def generate_recommendation(user_id):
    # Get the dominant VARK category for the user
    dominant_vark = calculate_vark_profile(user_id)

    # Get TOPSIS rankings for the user's responses
    user_responses = QuestionnaireResponse.query.filter_by(user_id=user_id).all()
    performance_scores = [[response.vark_score, response.knowledge_gap_score, response.time_spent_score] for response in user_responses]
    topsis_ranks = topsis_ranking(performance_scores)

    # Extract the TOPSIS rank for each question
    question_ranks = {response.question_id: rank for response, rank in zip(user_responses, topsis_ranks)}

    # Recommendation based on the dominant VARK category

    vark_recommendation = f"Based on your dominant VARK category ({dominant_vark}), consider exploring learning materials that align with this style."
    
    # Count the occurrences of V, A, R, and K
    vark_counts = Counter()
    for response in user_responses:
        vark_counts[response.answer_option_id] += 1
    
    ranked_vark = [vark for vark, _ in vark_counts.most_common()]
    print(ranked_vark)


    # Recommendation based on TOPSIS rankings
    topsis_recommendation = []
    for question_id, rank in sorted(question_ranks.items(), key=lambda x: x[1]):
        topsis_recommendation .append((rank, question_id))

    return vark_recommendation, topsis_recommendation, ranked_vark


def recommend_materials(vark_style):
    materials = {
        'V': 'Video tutorials and visual presentations and more movies',
        'A': 'Group discussions and study groups',
        'R': 'Written materials and textbooks',
        'K': 'Practical exercises and hands-on activities',
    }

    return materials.get(vark_style, 'No specific recommendation for this style.')



def populate_questions():
    vark_questions = [
        "When learning a new skill, do you prefer:\nV) Watching someone demonstrate it\nA) Listening to instructions or explanations\nR) Reading about it\nK) Trying it out yourself",
        "When studying for an exam, do you find it most effective to:\nR) Read and review your notes or textbooks\nV) Use visual aids like diagrams or charts\nA) Engage in group discussions or study groups\nK) Practice and apply the material through exercises or activities",
        "When receiving instructions, do you prefer them to be:\nR) Written or provided in text form\nV) Illustrated with diagrams or images\nA) Delivered verbally or explained through conversations\nK) Demonstrated through hands-on examples",
        "When trying to remember something, do you find it helpful to:\nR) Write it down or make notes\nV) Visualize or create mental images of the information\nA) Discuss or talk about it with others\nK) Engage in practical or hands-on activities related to the subject",
        "In a classroom or lectVisualize the problem or draw diagramure setting, do you learn best when:\nR) Reading the assigned materials beforehand\nV) Viewing slides or visual presentations\nA) Participating in discussions or asking questions\nK) Engaging in activities or experiments",
        "When conveying information to others, do you find it most effective to:\nR) Provide written instructions or handouts\nV) Use visual aids or presentations\nA) Engage in discussions or conversations\nK) Demonstrate or show examples",
        "When faced with a problem, do you prefer to:\nR) Analyze it by reading or researching\nV) Visualize possible solutions or draw diagrams\nA) Discuss it with others to gather different perspectives\nK) Experiment or try different approaches to solve it",
        "When exploring a new place or city, do you prefer to:\nR) Read guidebooks or maps\nV) Look at pictures or watch videos\nA) Interact with locals or ask for recommendations\nK) Explore and navigate the place on your own",
        "When learning a new language, do you find it helpful to:\nR) Read and study grammar rules\nV) Use visual aids like flashcards or pictorial representations\nA) Engage in conversations or language exchange with others\nK) Practice speaking and listening skills through immersion or real-life situations",
        "When trying to understand complex information, do you prefer to:\nR) Read detailed explanations or textbooks\nV) Use visualizations or diagrams to simplify the concepts\nA) Engage in discussions or debates to gain different viewpoints\nK) Break down the information into practical or tangible examples",
        "When memorizing information, do you find it effective to:\nR) Write and rewrite the information multiple times\nV) Create visual associations or mind maps\nA) Discuss and teach the material to someone else\nK) Actively practice and rehearse the information",
        "When learning a musical instrument, do you prefer to:\nR) Read sheet music or instructional books\nV) Watch videos or tutorials\nA) Take lessons or participate in group classes\nK) Practice and experiment with the instrument on your own",
        "When learning about historical events, do you find it most engaging to:\nR) Read books or articles about the subject\nV) Look at pictures or watch documentaries\nA) Engage in discussions or debates about the events\nK) Visit historical sites or museums related to the events",
        "When receiving feedback on your work, do you prefer it to be:\nR) Provided in written form with specific comments\nV) Presented visually with graphs or charts\nA) Delivered through verbal discussions or conversations\nK) Demonstrated through examples or models",
        "When solving math problems, do you find it helpful to:\nR) Read and analyze the problem step-by-step\nV) Visualize the problem or draw diagrams\nA) Discuss different problem-solving strategies with others\nK) Engage in hands-on practice or use manipulatives",
        "When learning about a new topic, do you prefer to:\nR) Read books or articles about the subject\nV) Watch videos or tutorials\nA) Engage in discussions or debates with others\nK) Explore and experiment with the subject firsthand",
    ]

    for question_str in vark_questions:
        question_data = parse_question(question_str)
        question = Question(question_text=question_data['question'])
        db.session.add(question)

        for option_data in question_data['options']:
            option = Option(question=question, option_id=option_data['id'], content=option_data['content'])
            db.session.add(option)

    db.session.commit()

def parse_question(question_str):
    parts = question_str.split('\n')
    question_text = parts[0]
    options = [{'id': part[0], 'content': part[3:]} for part in parts[1:] if part]
    
    # Convert numeric option IDs to corresponding VARK values
    for option in options:
        if option['id'] == '1':
            option['id'] = 'V'
        elif option['id'] == '2':
            option['id'] = 'A'
        elif option['id'] == '3':
            option['id'] = 'R'
        elif option['id'] == '4':
            option['id'] = 'K'

    return {'question': question_text, 'options': options}


def populate_hardware_software_memory_ports_questions():
    questions_data = [
        # Hardware
        ("Which component of a computer is responsible for storing data in the long term, even when the power is turned off?",
         ["a. RAM (Random Access Memory)", "b. CPU", "c. Hard drive", "d. Graphics card"], "Hardware"),

        ("What does GPU stand for in the context of computer hardware?",
         ["a. General Processing Unit", "b. Graphics Processing Unit", "c. General Peripheral Unit", "d. Global Processing Unit"],  "Hardware"),

        ("Which component of a computer system is responsible for managing communication between different hardware components?",
         ["a. Motherboard", "b. Hard drive", "c. Power supply", "d. RAM"], "Hardware"),
        ("Which of the following is a primary input device for a computer?",
         ["a. Monitor", "b. Keyboard", "c. Printer", "d. Speakers"],  "Hardware"),

        ("What is the purpose of the motherboard in a computer?",
         ["a. Provides power to the system", "b. Manages data storage", "c. Connects and facilitates communication between hardware components", "d. Executes software programs"],  "Hardware"),

        # Software
        ("Which type of software is designed to perform a specific task, such as word processing or spreadsheet calculations?",
         ["a. Operating system", "b. Application software", "c. System software", "d. Utility software"],"Software"),

        ("What is the purpose of an operating system (OS) on a computer?",
         ["a. Word processing", "b. Managing hardware and software resources", "c. Virus protection", "d. Graphic design"],"Software"),

        ("What type of software helps protect a computer from viruses, malware, and other security threats?",
         ["a. Word processor", "b. Firewall", "c. Antivirus software", "d. Spreadsheet software"],"Software"),

        ("Which programming language is often used for introductory programming courses for freshmen?",
         ["a. Java", "b. C++", "c. Python", "d. HTML"], "Software"),

        ("Which software is designed to manage and organize data in tabular form, often used for tasks such as budgeting and calculations?",
         ["a. Database software", "b. Presentation software", "c. Web browser", "d. Video editing software"],"Software"),

        # Memorye
        ("What is the primary function of RAM (Random Access Memory) in a computer system?",
         ["a. Long-term storage", "b. Temporary storage for running programs", "c. External data backup", "d. Graphic processing"],"Memory"),

        ("Which type of memory retains its content even when the power is turned off?",
         ["a. RAM", "b. ROM", "c. Cache memory", "d. Virtual memory"],"Memory"),

        ("Which of the following is a type of non-volatile memory used for long-term storage?",
         ["a. Cache memory", "b. RAM", "c. Flash memory", "d. Virtual memory"],"Memory"),

        ("What is the purpose of cache memory in a computer system?",
         ["a. Long-term storage", "b. Storing frequently accessed data for faster retrieval", "c. Permanent data retention", "d. Virtualization support"],"Memory"),

        ("Which type of memory is often used as a bridge between the CPU and RAM to enhance processing speed?",
         ["a. ROM", "b. Cache memory", "c. Flash memory", "d. Hard disk"],"Memory"),

        # Ports
        ("What type of port is typically used for connecting external hard drives and flash drives to a computer?",
         ["a. HDMI", "b. USB", "c. Thunderbolt", "d. Serial port"],"Ports"),

        ("Which port is commonly used for connecting a monitor to a computer for video output?",
         ["a. USB", "b. VGA", "c. Ethernet", "d. FireWire"],"Ports"),

        ("In networking, which port is associated with HTTP (Hypertext Transfer Protocol) traffic?",
         ["a. Port 80", "b. Port 443", "c. Port 21", "d. Port 25"],"Ports"),

        ("What is the primary function of an Ethernet port on a computer or network device?",
         ["a. Video output", "b. Data transfer over a wired network", "c. Power supply", "d. Audio output"],"Ports"),

        ("What is the purpose of an HDMI port on a TV or computer monitor?",
         ["a. Audio output", "b. Video output", "c. Data transfer", "d. Network connectivity"],"Ports"),
    ]

    for question_text, options, category in questions_data:
        question = QuestionMC(question_text=question_text, category=category)
        db.session.add(question)

        optionIDs = ['A', 'B', 'C', 'D']
        for i, option_text in enumerate(options):
            # Split option_text by dot
            option_parts = option_text.strip().split('. ')

            # Check if there are at least two elements after splitting
            if len(option_parts) >= 2:
                # Get the content after the dot and strip any leading/trailing whitespace
                option_content = option_parts[1].strip()

                # Use 'A', 'B', 'C', 'D' as option_id
                option = OptionMC(question=question, option_id=optionIDs[i], content=option_content)
                db.session.add(option)
            else:
                # Log or handle the case where the split result doesn't have enough elements
                print(f"Error: Unable to extract content from option: {option_text}")


    db.session.commit()



#pages

@app.route('/read/memory')
def read_memory_page():
    return render_template('Pages/Read_Memory.html')


@app.route('/hardware')
def hardware_page():
    return render_template('Pages/hardware_page.html')

@app.route('/software')
def software_page():
    return render_template('Pages/software_page.html')


@app.route('/ports')
def ports_page():
    return render_template('Pages/ports_page.html')


if __name__ == '__main__':
    with app.app_context():
        # Create the database tables before running the app
        db.create_all()
        # db.session.query(Question).delete()
        # db.session.query(Option).delete()
        # db.session.query(QuestionMC).delete()
        # db.session.query(OptionMC).delete()
        # db.session.query(QuestionnaireResponse).delete()
        # db.session.query(UserResponse).delete()
        db.session.commit()
        # populate_questions()
        # populate_hardware_software_memory_ports_questions()
        


    app.run(debug=True, host='0.0.0.0', port=5000)